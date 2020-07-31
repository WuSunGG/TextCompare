# coding=utf-8
import json
import re

import copy
from django.http import HttpResponse
from django.shortcuts import render

from simhash.Response import buildResponseError, buildResponseSuccess
from simhash.simhash import SimHash
import difflib
from dev_SimHash import SimHash as SimHash2, parseToEchartsData, generateHTML, getPaperSimhash, comparePaper
from dev_SimHash import getPapersResults
import uuid

from simhash.templatetags.myfilter import stringtofile


def index(request):
    context = {
        'answer': '1',
        'tx1': '1',
        'tx2': '1',
        'issame': "",  # 是否显示
        'hmdistance': '',  # 海明距离,
        'dedistance': '',  # 判断距离,
        'diffe': ''  # 对比结果
    }

    try:
        simhash = SimHash()
        tx1 = (request.POST['tx1'])
        tx2 = (request.POST['tx2'])
        context['tx1'] = tx1
        context['tx2'] = tx2
        if len(tx1) <= 2 or len(tx2) <= 2:
            context['answer'] = "对比文本不能小于2字符"
        else:
            hash1 = simhash.hash(tx1)
            hash2 = simhash.hash(tx2)
            distince = simhash.getDistince(hash1, hash2)
            value = 20
            context['hmdistance'] = distince
            context['dedistance'] = value
            context['issame'] = distince <= value
            res = "海明距离：" + str(distince) + " 判定距离：" + str(value) + " 是否相似 " + str(distince <= value)
            context['answer'] = res

            dtx1 = tx1.splitlines()
            dtx2 = tx2.splitlines()
            d = difflib.HtmlDiff()
            dres = d.make_file(dtx1, dtx2, charset="utf-8")
            context['diffe'] = dres
    except KeyError:
        context['answer'] = '空字符串'
        return render(request, 'simhash.html', context)
    return render(request, 'simhash.html', context)


def index2(request):
    res = getPapersResults()
    context = {
        'results': res,
        'dec_distance': 20
    }

    return render(request, 'papercompare.html', context)


def index21(request, distance=20):
    res = getPapersResults(distance)
    context = {
        'results': res,
        'dec_distance': distance
    }
    return render(request, 'papercompare.html', context)


def index3(request, distance=20):
    rhtml = []
    cres = getPapersResults(distance)
    for stuCompare in cres:
        t_html = {}
        data, links = parseToEchartsData(stuCompare, cres[stuCompare])
        data_json = json.dumps(data)
        links_json = json.dumps(links)
        html = generateHTML(data_json, links_json, uuid.uuid1())
        t_html['name'] = stuCompare
        t_html['html'] = html
        rhtml.append(t_html)

    context = {
        'results': rhtml,
    }

    return render(request, 'papercompare_pic.html', context)


def htmlcheckv1(request):
    print(request.body)
    try:
        req = json.loads(request.body)
        req['type']
        req['papers']
        req['distance']
    except Exception as e:
        return buildResponseError("request format is error : {}".format(e))

    try:
        papersObj = {}
        for paper in req['papers']:
            tsimhash = getPaperSimhash(paper['content'])
            # print(t_file)
            papersObj[paper['sno']] = paper
            print ( papersObj[paper['sno']])
            papersObj[paper['sno']]['content'] = paper['content']
            papersObj[paper['sno']]['simhash'] = tsimhash
            # print(targets)

            # 计算每篇文章和其他文章的相似度
        cres = comparePaper(papersObj, req['distance'])
        print (cres)
        return buildResponseSuccess(cres)

    except Exception as e:
        return buildResponseError("paper format is error: {}".format(e))


from docx import Document


def compare(request, fp="", tp=""):
    context = {}
    document = Document("./works/15软工-37-郑梦月.docx")
    tx1 = ""
    for paragraph in document.paragraphs:
        tx1 += paragraph.text + "\n"

    document = Document("./works/15软工-38-臧庆照.docx")
    tx2 = ""
    for paragraph in document.paragraphs:
        tx2 += paragraph.text + "\n"

    # dtx1=tx1.splitlines()
    # dtx2=tx2.splitlines()

    dtx1 = re.split('(。|！|\!|\.|？|\?|\n|，)', tx1)  # 保留分割符
    dtx2 = re.split('(。|！|\!|\.|？|\?|\n|，)', tx2)  # 保留分割符
    #
    new_sent1s = []
    for i in range(int(len(dtx1) / 2)):
        sent = dtx1[2 * i] + dtx1[2 * i + 1]
        new_sent1s.append(sent)

    new_sent2s = []
    for i in range(int(len(dtx2) / 2)):
        sent = dtx2[2 * i] + dtx2[2 * i + 1]
        new_sent2s.append(sent)

    d = difflib.HtmlDiff()
    dres = d.make_file(new_sent1s, new_sent2s, charset="utf-8")
    context['diffe'] = dres
    print(context)
    return render(request, 'compare.html', context)
