# coding=utf-8
# simhash  文本相似度计算方法
# 原文链接：https://blog.csdn.net/al_xin/java/article/details/38919361
# C++ 实现: https://github.com/yanyiwu/simhash
# Python实现：https://leons.im/posts/a-python-implementation-of-simhash-algorithm/
# 算法主要原理：
# 1、分词，把需要判断文本分词形成这个文章的特征单词。最后形成去掉噪音词的单词序列并为每个词加上权重，我们假设权重分为5个级别（1~5）。比如：“ 美国“51区”雇员称内部有9架飞碟，曾看见灰色外星人 ” ==> 分词后为 “ 美国（4） 51区（5） 雇员（3） 称（1） 内部（2） 有（1） 9架（3） 飞碟（5） 曾（1） 看见（3） 灰色（4） 外星人（5）”，括号里是代表单词在整个句子里重要程度，数字越大越重要。
# 2、hash，通过hash算法把每个词变成hash值，比如“美国”通过hash算法计算为 100101,“51区”通过hash算法计算为 101011。这样我们的字符串就变成了一串串数字，还记得文章开头说过的吗，要把文章变为数字计算才能提高相似度计算性能，现在是降维过程进行时。
# 3、加权，通过 2步骤的hash生成结果，需要按照单词的权重形成加权数字串，比如“美国”的hash值为“100101”，通过加权计算为“4 -4 -4 4 -4 4”；“51区”的hash值为“101011”，通过加权计算为 “ 5 -5 5 -5 5 5”。
# 4、合并，把上面各个单词算出来的序列值累加，变成只有一个序列串。比如 “美国”的 “4 -4 -4 4 -4 4”，“51区”的 “ 5 -5 5 -5 5 5”， 把每一位进行累加， “4+5 -4+-5 -4+5 4+-5 -4+5 4+5” ==》 “9 -9 1 -1 1 9”。这里作为示例只算了两个单词的，真实计算需要把所有单词的序列串累加。
# 5、降维，把4步算出来的 “9 -9 1 -1 1 9” 变成 0 1 串，形成我们最终的simhash签名。 如果每一位大于0 记为 1，小于0 记为 0。最后算出结果为：“1 0 1 0 1 1”。
# demo
# “你妈妈喊你回家吃饭哦，回家罗回家罗” 和 “你妈妈叫你回家吃饭啦，回家罗回家罗”。
# 通过simhash计算结果为：
# 1000010010101101111111100000101011010001001111100001001011001011
# 1000010010101101011111100000101011010001001111100001101010001011
import base64
import json
import math
import os
import jieba
import jieba.analyse
from docx import Document
from string import Template


class SimHash(object):
    def __init__(self):
        pass

    def getBinStr(self, source):
        """
        将 原文本 转换为 二进制
        如 资源 转换为 0100000000111101100101001111110101101000101000010111010001011110
        :param source: 原文本
        :return:  二进制
        """
        if source == "":
            return str(0)
        else:
            x = ord(source[0]) << 7
            m = 1000003
            mask = 2 ** 128 - 1
            for c in source:
                x = ((x * m) ^ ord(c)) & mask
            x ^= len(source)
            if x == -1:
                x = -2
            x = bin(x).replace('0b', '').zfill(64)[-64:]
            return str(x)

    def getWeight(self, source):
        # fake weight with keyword
        return ord(source)

    def unwrap_weight(self, arr):
        ret = ""
        for item in arr:
            tmp = 0
            if int(item) > 0:
                tmp = 1
            ret += str(tmp)
        return ret

    def hash(self, rawstr):
        """

        :param rawstr: 文本
        :return:
        """
        # 1 分词
        seg = jieba.cut(rawstr)
        # print('/'.join(seg))
        # 结果：州/道/教育/科技/有限公司/（/ARKTAO/）/的/一员/是/孙武/同学/。

        # 文字关键词
        #  找到文章的关键词，词频解析为 （词语，权重）,如
        #  {tuple: 2} ('资源', 0.26808614843899997)
        #  {tuple: 2} ('系统', 0.21585966056762962)
        #  {tuple: 2} ('用户', 0.14529049043116665)
        keywords = jieba.analyse.extract_tags("|".join(seg), topK=1000, withWeight=True)
        # 结果
        #         0 = {tuple: 2} ('ARKTAO', 1.7078239289857142)
        #         1 = {tuple: 2} ('孙武', 1.426950663792857)
        #         2 = {tuple: 2} ('一员', 1.1383474949328571)
        #         3 = {tuple: 2} ('同学', 1.0294450847785714)
        #         4 = {tuple: 2} ('科技', 0.8331820990585713)
        #         5 = {tuple: 2} ('教育', 0.8111989397285715)
        #         6 = {tuple: 2} ('有限公司', 0.7419558412757142)

        ret = []
        # keyword = 'ARKTAO' weight = 1.7078239289857142
        for keyword, weight in keywords:
            # 2 hash # ARKTAO => '1011000111000111010001010101110011100101010011101010001011010110'
            binstr = self.getBinStr(keyword)  # '0100000000111101100101001111110101101000101000010111010001011110'

            keylist = []
            # res
            # [2, -2, 2, 2, -2, -2, -2, 2, 2, 2, -2, -2, -2, 2, 2, 2, -2, 2, -2, -2, -2, 2, -2, 2, -2, 2, -2, 2, 2, 2, -2, -2, 2, 2, 2, -2, -2, 2, -2, 2, -2, 2, -2, -2, 2, 2, 2, -2, 2, -2, 2, -2, -2, -2, 2, -2, 2, 2, -2, 2, -2, 2, 2, -2]
            # [1, -1, 1, -1, -1, 1, -1, 1, 1, -1, 1, 1, 1, 1, 1, -1, -1, -1, 1, 1, -1, 1, 1, 1, -1, -1, 1, -1, 1, 1, 1, 1, -1, 1, 1, -1, 1, -1, -1, 1, -1, 1, -1, 1, -1, 1, 1, 1, 1, -1, 1, 1, -1, 1, 1, -1, 1, -1, -1, 1, 1, -1, 1, 1]
            # 3 加权
            for c in binstr:  # c 是 binstr 中的值，二进制，1或者0
                weight = math.ceil(weight)
                # 4 合并
                if c == "1":
                    keylist.append(int(weight))
                else:
                    keylist.append(-int(weight))

            ret.append(keylist)
            # ret = {list: 7} [[2, -2, 2, 2, -2, -2, -2, 2, 2, 2, -2, -2, -2, 2, 2, 2, -2, 2, -2, -2, -2, 2, -2, 2, -2, 2, -2, 2, 2, 2, -2, -2, 2, 2, 2, -2, -2, 2, -2, 2, -2, 2, -2, -2, 2, 2, 2, -2, 2, -2, 2, -2, -2, -2, 2, -2, 2, 2, -2, 2, -2, 2, 2, -2], [-2, -2, 2, -2, 2, -2, -2, 2, 2, -2, -2, -2, 2, -2, 2, -2, -2, 2, -2, 2, -2, 2, -2, -2, 2, -2, 2, -2, -2, 2, 2, 2, 2, 2, 2, -2, 2, 2, 2, -2, 2, -2, -2, 2, -2, -2, 2, 2, -2, 2, -2, -2, 2, 2, 2, 2, 2, -2, 2, -2, 2, 2, 2, 2], [-2, -2, 2, -2, -2, -2, 2, 2, -2, 2, 2, 2, 2, -2, -2, -2, -2, 2, 2, 2, -2, -2, 2, 2, -2, -2, -2, 2, 2, 2, -2, -2, 2, -2, 2, -2, -2, 2, 2, -2, -2, -2, -2, -2, 2, 2, 2, 2, -2, -2, 2, 2, 2, 2, 2, -2, -2, 2, -2, 2, 2, -2, 2, -2], [-2, -2, 2, -2, -2, 2, 2, -2, -2, -2, 2, 2, 2, -2, -2, -2, -2, 2, 2, -2, -2, 2, 2, 2, -2, -2, -2, 2, 2, -2, 2, -2, 2, 2, 2, 2, -2, 2, 2, -2, -2, 2, -2, 2, -2, 2, 2, 2, 2, -2, 2, 2, -2, 2, 2, -2, -2, 2, -2, -2, -2, -2, -2, -2], [-1, -1, 1, 1, -1, 1, 1, 1, -1, 1, 1, -1, -1, 1, -1, 1, -1, 1, 1, -1, -1, -1, -1, 1, 1, -1, -1, 1,...
            #  0 = {list: 64} [2, -2, 2, 2, -2, -2, -2, 2, 2, 2, -2, -2, -2, 2, 2, 2, -2, 2, -2, -2, -2, 2, -2, 2, -2, 2, -2, 2, 2, 2, -2, -2, 2, 2, 2, -2, -2, 2, -2, 2, -2, 2, -2, -2, 2, 2, 2, -2, 2, -2, 2, -2, -2, -2, 2, -2, 2, 2, -2, 2, -2, 2, 2, -2]
            #  1 = {list: 64} [-2, -2, 2, -2, 2, -2, -2, 2, 2, -2, -2, -2, 2, -2, 2, -2, -2, 2, -2, 2, -2, 2, -2, -2, 2, -2, 2, -2, -2, 2, 2, 2, 2, 2, 2, -2, 2, 2, 2, -2, 2, -2, -2, 2, -2, -2, 2, 2, -2, 2, -2, -2, 2, 2, 2, 2, 2, -2, 2, -2, 2, 2, 2, 2]
            #  2 = {list: 64} [-2, -2, 2, -2, -2, -2, 2, 2, -2, 2, 2, 2, 2, -2, -2, -2, -2, 2, 2, 2, -2, -2, 2, 2, -2, -2, -2, 2, 2, 2, -2, -2, 2, -2, 2, -2, -2, 2, 2, -2, -2, -2, -2, -2, 2, 2, 2, 2, -2, -2, 2, 2, 2, 2, 2, -2, -2, 2, -2, 2, 2, -2, 2, -2]
            #  3 = {list: 64} [-2, -2, 2, -2, -2, 2, 2, -2, -2, -2, 2, 2, 2, -2, -2, -2, -2, 2, 2, -2, -2, 2, 2, 2, -2, -2, -2, 2, 2, -2, 2, -2, 2, 2, 2, 2, -2, 2, 2, -2, -2, 2, -2, 2, -2, 2, 2, 2, 2, -2, 2, 2, -2, 2, 2, -2, -2, 2, -2, -2, -2, -2, -2, -2]
            #  4 = {list: 64} [-1, -1, 1, 1, -1, 1, 1, 1, -1, 1, 1, -1, -1, 1, -1, 1, -1, 1, 1, -1, -1, -1, -1, 1, 1, -1, -1, 1, -1, 1, -1, -1, 1, -1, -1, 1, -1, -1, 1, -1, 1, -1, 1, -1, -1, -1, -1, -1, 1, -1, 1, 1, -1, 1, -1, 1, 1, -1, 1, 1, -1, -1, -1, 1]
            #  5 = {list: 64} [-1, -1, 1, -1, 1, 1, 1, -1, -1, -1, -1, 1, -1, 1, 1, -1, -1, 1, 1, 1, 1, 1, -1, 1, 1, -1, 1, -1, -1, 1, 1, 1, -1, 1, -1, 1, -1, -1, -1, -1, -1, -1, 1, 1, -1, -1, 1, 1, -1, 1, 1, -1, 1, -1, 1, -1, -1, 1, 1, 1, 1, -1, 1, 1]
            #  6 = {list: 64} [1, -1, 1, -1, -1, 1, -1, 1, 1, -1, 1, 1, 1, 1, 1, -1, -1, -1, 1, 1, -1, 1, 1, 1, -1, -1, 1, -1, 1, 1, 1, 1, -1, 1, 1, -1, 1, -1, -1, 1, -1, 1, -1, 1, -1, 1, 1, 1, 1, -1, 1, 1, -1, 1, 1, -1, 1, -1, -1, 1, 1, -1, 1, 1]
        # 5 降维
        rows = len(ret)
        cols = len(ret[0])
        result = []
        for i in range(cols):
            tmp = 0
            for j in range(rows):
                tmp += int(ret[j][i])
            if tmp > 0:
                tmp = "1"
            elif tmp <= 0:
                tmp = "0"
            result.append(tmp)
        return "".join(result)

    def getDistince(self, hashstr1, hashstr2):
        """
        计算距离
        :param hashstr1: 第一个 simHash
        :param hashstr2: 第二个 simHash
        :return:
        """
        length = 0
        for index, char in enumerate(hashstr1):
            if char == hashstr2[index]:
                continue
            else:
                length += 1
        return length


def debug(msg, msg2=""):
    print("debug>> " + msg, msg2)


def comparePaper(papersObj,decision_distince=20):
    """
    结果：
    {
        stuwork:[
            distance:{name,simhash} 排好序的。0表示完全一样，36表示完全不一样，越大越不相似，越小越相似
        ]
    }

    :param papersObj:
    :param simhash:
    :return:
    """
    simhash=SimHash()
    res = {}
    for paper in papersObj:
        res[paper] = []
        # print(paper)
        # ./works/15软工-44-张大发.docx
        current_hash = papersObj[paper]['simhash']
        for comPaper in papersObj:
            if comPaper == paper:
                t_res = papersObj[comPaper].copy()
                t_res['distance'] = -1
                del t_res['simhash']
                del t_res['content']
                res[paper].append(t_res)
                continue
            distance = simhash.getDistince(current_hash, papersObj[comPaper]['simhash'])
            if distance <= decision_distince:
                t_res = papersObj[comPaper].copy()
                t_res['distance'] = distance
                del t_res['simhash']
                del t_res['content']
                t_res['percent']=(36-distance)/36 *100
                # del t_res['simhash']
                res[paper].append(t_res)

        res[paper].sort(key=takeSecond)
    return res


def takeSecond(elem):
    return elem["distance"]


def getStuNameFromFileName(str):
    if len(str) <= 0: return ""
    sindex = str.rfind("-")
    endindex = str.rfind(".")
    # print(sindex, endindex)
    return (str[sindex + 1:endindex])


def parseToEchartsData(key, values):
    """
    单个学生的结果变为 json
    data: [{ // data
             symbolSize: 100, // 决定当前图形的大小
             name: 'sunwu', //名字，用来做连线的
             x: 300, // x 坐标
             y: 300 // y 坐标
         }],
         links: [{ // 关系
             source: 'sunwu',
             target: 'lbn'
         }],
    :param key: "./works/15软工-15-陈鸿.doc"
    :param value:
    :return:
    """

    datas = []
    links = []
    x = 800
    y = 50
    if len(values) >= 1:
        stuName = getStuNameFromFileName(key)
        # datas.append({"name": stuName, "x": 50, "y": 300 * len(values) / 2, "symbolSize": 50})
        datas.append({"name": stuName,  "symbolSize": 50})
        for apd in values:
            tname = getStuNameFromFileName(apd['name'])
            # datas.append({"name": tname, "x": x, "y": y, "symbolSize": (36 - apd['distance']) * 2})
            datas.append({"name": tname, "symbolSize": (36 - apd['distance']) * 5})
            # links.append({"source": stuName, "target":tname,"lineStyle":{"curveness":0.5,"width":36-apd['distance']}})
            # links.append({"source": stuName, "target": tname, "lineStyle": {"curveness":0.2,"width": 36 - apd['distance']}})
            links.append({"source": stuName, "target": tname, "lineStyle": {"curveness":0.2}})
            y += 300

    return datas, links


def generateHTML(data_json, links_json,id='main'):
    tpl = Template("""
     <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>ECharts</title>
        <!-- 引入 echarts.js -->
       <script src="https://cdn.bootcdn.net/ajax/libs/echarts/4.8.0/echarts.min.js"></script>
    </head>
    <body >
        <!-- 为ECharts准备一个具备大小（宽高）的Dom -->
        <div id="${id}" style="width:1280px;height:720px;"></div>
        <script type="text/javascript">
            // 基于准备好的dom，初始化echarts实例
            var myChart = echarts.init(document.getElementById('${id}'));


            var data= ${data}  ;
            var links= ${links} ;


                      // 指定图表的配置项和数据
            var option = {
                title: {
                    text: '关系图（圆越大，抄袭概率越大）'
                },
                tooltip: {},
                animationDurationUpdate: 1,
                 focusNodeAdjacency: true,
                animationEasingUpdate: 'quinticInOut',
                series: [
                    {
                        type: 'graph',
                        layout: 'force',
                        roam: true,
                        symbolSize: 50,// 决定全局图像的大小，如圆
                        label: {
                            show: true
                        },
                        edgeSymbol: ['rect', 'arrow'],
                        edgeSymbolSize: [4, 10],
                        edgeLabel: {
                            fontSize: 10
                        },
                        data: data,
                        links:links,
                        lineStyle: {
                            opacity: 0.9,
                            width: 1,
                            curveness: 0.1
                        },
                        force: {
                            repulsion: 2000
                        }
                    }
                ]
            };
            // 使用刚指定的配置项和数据显示图表。
            myChart.setOption(option);
        </script>
    </body>
    </html>
    """)
    #  var data= ${data}  ;
    #             var links= ${links} ;
    return  tpl.substitute(data=data_json, links=links_json,id=id)


def getPaperSimhash(paper):
    """
    获取文章的 simhash 值
    :param paper:
    :return:
    """
    simhash = SimHash()
    return simhash.hash(paper)

def getPapersResults(distance=20):
    simhash = SimHash()
    workpath = "./works"

    #  计算文章的 simHash
    # {
    #   filename:{
    #       content:""
    #       simhash:""
    #   }
    # }
    papersObj = {}
    for file in os.listdir(workpath):
        t_file = os.path.join(workpath, file)
        # print(t_file)
        #     ./works/15软工-15-陈鸿.doc
        # ./works/15软工-17-胡应学.docx

        # 过滤掉 .doc 文件
        if t_file.find(".docx") > -1:
            # print(t_file)
            papersObj[t_file] = {}
            path = t_file
            document = Document(path)
            tpaperTxt = ""
            for paragraph in document.paragraphs:
                tpaperTxt += paragraph.text + "/r/n"
            papersObj[t_file]['content'] = tpaperTxt
            papersObj[t_file]['simhash'] = simhash.hash(tpaperTxt)
        # print(targets)

    # 计算每篇文章和其他文章的相似度
    cres = comparePaper(papersObj,distance)

    return cres
    # # 数据处理为 echarts
    # for stuCompare in cres:
    #     data, links = parseToEchartsData(stuCompare, cres[stuCompare])
    #     data_json = json.dumps(data)
    #     links_json = json.dumps(links)
    #
    #     html = generateHTML(data_json, links_json)
    #     filename = 'a.html'
    #     with open(filename, 'w') as f:
    #         f.write(html)
    #     print("写入完成")
    #     exit(0)


if __name__ == "__main__":
    # word 转 txt
    # https://python-docx.readthedocs.io/en/latest/index.html
    # pip install python-docx



    simhash = SimHash()

    workpath = "./works"

    #  计算文章的 simHash
    # {
    #   filename:{
    #       content:""
    #       simhash:""
    #   }
    # }
    papersObj = {}
    for file in os.listdir(workpath):
        t_file = os.path.join(workpath, file)
        # print(t_file)
        #     ./works/15软工-15-陈鸿.doc
        # ./works/15软工-17-胡应学.docx

        # 过滤掉 .doc 文件
        if t_file.find(".docx") > -1:
            # print(t_file)
            papersObj[t_file] = {}
            path = t_file
            document = Document(path)
            tpaperTxt = ""
            for paragraph in document.paragraphs:
                tpaperTxt += paragraph.text + "/r/n"
            papersObj[t_file]['content'] = tpaperTxt
            papersObj[t_file]['simhash'] = simhash.hash(tpaperTxt)
        # print(targets)

    # 计算每篇文章和其他文章的相似度
    cres = comparePaper(papersObj,20)
    # './works/15软工-45-吴强.docx' = {list: 5} [{'distance': 13, 'name': './works/15软工-46-罗永双.docx'}, {'distance': 14, 'name': './works/15软工-36-刘思乾.docx'}, {'distance': 15, 'name': './works/15软工-35-张德志.docx'}, {'distance': 15, 'name': './works/15软工-37-郑梦月.docx'}, {'distance': 15, 'name': './works/15软工-38-臧庆照.docx'}]
    #  0 = {dict: 2} {'distance': 13, 'name': './works/15软工-46-罗永双.docx'}
    # './works/15软工-46-罗永双.docx' = {list: 3} [{'distance': 13, 'name': './works/15软工-45-吴强.docx'}, {'distance': 14, 'name': './works/15软工-38-臧庆照.docx'}, {'distance': 15, 'name': './works/15软工-30-詹源.docx'}]
    # './works/15软工-49-杨豪.docx' = {list: 0} []
    # './works/15软工-29-胡佳乐.docx' = {list: 0} []
    # './works/15软工-30-詹源.docx' = {list: 1} [{'distance': 15, 'name': './works/15软工-46-罗永双.docx'}]
    # 数据处理为 echarts
    for stuCompare in cres:
        data, links = parseToEchartsData(stuCompare, cres[stuCompare])
        data_json = json.dumps(data)
        links_json = json.dumps(links)

        html=generateHTML(data_json,links_json)
        filename = 'test/a.html'
        with open(filename, 'w') as f:
            f.write(html)
        print("写入完成")
        exit(0)


    # print(cres)
    # cres = {dict: 41} {'./works/15-软工-47-杨宗流.docx':
    #  './works/15-软工-47-杨宗流.docx' = {list: 19} [
    #   00 = {dict: 2} {'distance': 13, 'name': './works/15软工-11-曾祥武.docx'}
    #   01 = {dict: 2} {'distance': 15, 'name': './works/15软工-45-吴强.docx'}
    #   02 = {dict: 2} {'distance': 16, 'name': './works/15软工-25-胡洋洋.docx'}
    #   03 = {dict: 2} {'distance': 16, 'name': './works/15软工-29-胡佳乐.docx'}
    #   04 = {dict: 2} {'distance': 17, 'name': './works/15软工-08-孔莉.docx'}
    #   05 = {dict: 2} {'distance': 17, 'name': './works/15软工-23-尹涵.docx'}
    #   06 = {dict: 2} {'distance': 17, 'name': './works/15软工-38-臧庆照.docx'}
    #   07 = {dict: 2} {'distance': 19, 'name': './works/15软工-09-朱鹏成.docx'}
    #   08 = {dict: 2} {'distance': 19, 'name': './works/15软工-10-鲜飞.docx'}
    #   09 = {dict: 2} {'distance': 19, 'name': './works/15软工-27-陈涛.docx'}
    #   10 = {dict: 2} {'distance': 19, 'name': './works/15软工-32-王凡.docx'}
    #   11 = {dict: 2} {'distance': 19, 'name': './works/15软工-36-刘思乾.docx'}
    #   12 = {dict: 2} {'distance': 19, 'name': './works/15软工-41-陈隐农.docx'}
    #   13 = {dict: 2} {'distance': 19, 'name': './works/15软工-46-罗永双.docx'}
    #   14 = {dict: 2} {'distance': 20, 'name': './works/15软工-37-郑梦月.docx'}
    #   15 = {dict: 2} {'distance': 20, 'name': './works/15软工-39-张宇帆.docx'}
    #   16 = {dict: 2} {'distance': 20, 'name': './works/15软工-40-王鼎燊.docx'}
    #   17 = {dict: 2} {'distance': 20, 'name': './works/15软工-42-朱仕全.docx'}
    #   18 = {dict: 2} {'distance': 20, 'name': './works/15软工-43-李秋美.docx'}
